from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document


WORKDIR = Path(r"C:\Users\SummerNeal\OneDrive\OneDrive - True North Companies, LLC\Documents\Agent-Identity")
DOCX_PATH = WORKDIR / "mcp_overprivileged_demo_guide_polished.docx"


PARA_REPLACEMENTS = {
    "DEMO GUIDE": "Demo Guide",
    "Step-by-step build guide for security blog demonstration": "Step-by-step guide for a security blog demonstration",
    "This demo shows how an MCP server configured with excessive permissions leaks sensitive data through a Copilot Studio agent — without any malicious payload or prompt injection. The risk is structural: the agent faithfully exposes whatever the MCP server has access to, simply by responding to normal user questions.": "This demo shows how an MCP server with excessive permissions can leak sensitive data through a Copilot Studio agent, even when the user asks only ordinary business questions. The risk is structural: if the server can access sensitive data, the agent may faithfully surface that data in its response without any prompt injection or overtly malicious input.",
    "The scenario:": "In this scenario:",
    "A citizen developer builds an agent to answer HR questions": "A citizen developer builds an agent to answer HR questions.",
    "They connect an MCP server that has broad filesystem or environment access": "They connect an MCP server with broad file system or environment access.",
    "A user asks an innocent question: \"What is our leave policy?\"": "A user asks an ordinary question: \"What is our leave policy?\"",
    "The MCP server response includes secrets, keys, or internal paths alongside the answer": "The MCP server response includes secrets, keys, or internal paths alongside the answer.",
    "The agent surfaces that data to the user verbatim": "The agent returns that data to the user verbatim.",
    "No alert fires. The KQL query detects the MCP tool exists — but not what it returned": "No content-level alert necessarily fires. The KQL query can show that the MCP tool exists, but not what it returned during the conversation.",
    "What you will build": "What You Will Build",
    "Three components work together:": "Three components work together in this demo:",
    "The vulnerability demonstrated: the MCP server is given access it does not need. When the agent calls the tool to answer a benign question, the server returns sensitive data alongside legitimate content. The agent has no way to distinguish the two and surfaces everything.": "The vulnerability is simple: the MCP server is granted access it does not need. When the agent calls the tool to answer a benign question, the server returns sensitive data alongside legitimate content, and the agent has no reliable way to separate the two before surfacing the response.",
    "Python environment": "Python Environment",
    "Go to https://ngrok.com and sign up for a free account": "Go to https://ngrok.com and sign up for a free account.",
    "Download and install ngrok for your platform": "Download and install ngrok for your platform.",
    "Create a file called mcp_overprivileged.py with the following content:": "Create a file named mcp_overprivileged.py with the following content:",
    "Run the server": "Run the Server",
    "Expose via ngrok": "Expose the Server with ngrok",
    "Configure the agent instructions": "Configure the Agent Instructions",
    "In the Instructions panel, replace the default text with:": "In the Instructions panel, replace the default text with the following:",
    "Navigate to https://copilotstudio.microsoft.com": "Navigate to https://copilotstudio.microsoft.com.",
    "Sign in with your demo tenant account": "Sign in with your demo tenant account.",
    "Click \"Create\" in the left sidebar, then \"New agent\"": "Click \"Create\" in the left sidebar, then \"New agent.\"",
    "Name the agent: \"HR Assistant Demo\"": "Name the agent: \"HR Assistant Demo.\"",
    "In the Description field, enter: \"Answers employee HR questions about leave, benefits, and expenses\"": "In the Description field, enter: \"Answers employee HR questions about leave, benefits, and expenses.\"",
    "Click Create": "Click Create.",
    "Inside your agent, click the Tools tab (or Actions, depending on your Copilot Studio version)": "Inside your agent, open the Tools tab (or Actions, depending on your Copilot Studio version).",
    "Click \"Add a tool\" or \"+ Add action\"": "Click \"Add a tool\" or \"+ Add action.\"",
    "Select \"Model Context Protocol (MCP)\" from the connector type list": "Select \"Model Context Protocol (MCP)\" from the connector list.",
    "Click \"Connect\" — Copilot Studio will discover the available tools from your server": "Click \"Connect.\" Copilot Studio will discover the available tools from your server.",
    "You should see two tools appear: get_hr_policy and get_employee_count": "You should see two tools: get_hr_policy and get_employee_count.",
    "Enable both tools": "Enable both tools.",
    "Click \"Save\" to confirm the MCP connection": "Click \"Save\" to confirm the MCP connection.",
    "Open the Test panel in Copilot Studio (right side of the screen). Send these prompts in order:": "Open the Test panel in Copilot Studio (right side of the screen) and send these prompts in order:",
    "Prompt 1 — benign HR question": "Prompt 1 - Benign HR Question",
    "What the agent should return:": "Expected result:",
    "The legitimate leave policy answer": "The legitimate leave-policy answer.",
    "Followed by the DEBUG block showing the FAKE_SECRETS values": "A DEBUG block showing the FAKE_SECRETS values.",
    "Any matching environment variables from the server process": "Any matching environment variables from the server process.",
    "Prompt 2 — department headcount": "Prompt 2 - Department Headcount",
    "The headcount: 142": "The headcount: 142.",
    "Followed by the DB connection string in the DEBUG block": "A DEBUG block showing the database connection string.",
    "Prompt 3 — ask it to summarise everything it knows": "Prompt 3 - Ask the Agent to Summarize Everything It Knows",
    "This prompt often causes the agent to enumerate everything in its context window — including the leaked credential strings from earlier tool calls.": "This prompt often causes the agent to enumerate everything in its context window, including credential strings leaked by earlier tool calls.",
    "Open Microsoft Sentinel (or Defender XDR Advanced Hunting) and run the following query to show that the MCP tool configuration is detectable:": "Open Microsoft Sentinel (or Defender XDR Advanced Hunting) and run the following query to show that the MCP tool configuration is detectable:",
    "Your HR Assistant Demo agent should appear in the results, with MCPTools showing [\"hr-assistant-mcp\"].": "Your HR Assistant Demo agent should appear in the results, with MCPTools showing [\"hr-assistant-mcp\"] if the connection has been saved and telemetry is available.",
    "The detection gap — make this point in your blog": "Detection Gap - Make This Point in Your Blog",
    "The KQL query tells you the MCP tool EXISTS on the agent. It does NOT tell you:": "The KQL query shows that the agent has an MCP tool. It does not tell you:",
    "These are the key screenshots to capture for maximum impact:": "Capture these screenshots for the strongest blog visuals:",
    "Recommended mitigations": "Recommended Mitigations",
    "Include these in your blog as the defensive takeaways:": "Include these defensive takeaways in your blog:",
    "After the demo, clean up to avoid leaving an open tunnel or misconfigured agent:": "After the demo, clean up to avoid leaving an open tunnel or a misconfigured agent behind:",
    "Stop the ngrok tunnel (Ctrl+C in terminal 2)": "Stop the ngrok tunnel (Ctrl+C in terminal 2).",
    "Stop the Python MCP server (Ctrl+C in terminal 1)": "Stop the Python MCP server (Ctrl+C in terminal 1).",
    "In Copilot Studio, open the agent and remove the MCP tool under Tools": "In Copilot Studio, open the agent and remove the MCP tool under Tools.",
    "Set the agent status to \"Unpublished\" or delete it": "Set the agent status to \"Unpublished\" or delete it.",
    "Verify the KQL query no longer returns this agent (or shows AgentStatus = Deleted)": "Verify that the KQL query no longer returns this agent, or that it shows AgentStatus = Deleted.",
    "Built for security blog demonstration purposes. All credential strings in this guide are non-functional placeholders.": "Built for security blog demonstration purposes. All credential strings in this guide are nonfunctional placeholders.",
}


CELL_REPLACEMENTS = {
    "Demo complexity\nLow": "Complexity\nLow",
    "Time to build\n~30 min": "Build time\n~30 minutes",
    "Risk illustrated\nData leakage via MCP": "Risk illustrated\nData leakage through MCP",
    "This guide is for controlled security demonstration and blog content only. Build in a dedicated demo tenant. Never run against production data or user accounts.": "This guide is intended only for a controlled security demonstration and related blog content. Build it in a dedicated demo tenant, and never run it against production data or user accounts.",
    "Developer or demo tenant — never production": "Use a developer or demo tenant - never production.",
    "Trial or paid — Power Platform environment needed": "Trial or paid; requires a Power Platform environment.",
    "For the local MCP server (FastMCP)": "Required for the local FastMCP server.",
    "Exposes localhost to Copilot Studio over HTTPS": "Exposes localhost to Copilot Studio over HTTPS.",
    "macOS, Linux, or Windows WSL2": "macOS, Linux, or Windows with WSL2.",
    "A free Microsoft 365 Developer tenant (via Microsoft 365 Developer Program) is ideal. It gives you E5 licences at no cost for 90 days.": "A Microsoft 365 Developer tenant is often the easiest option for this demo. Review the current Microsoft 365 Developer Program terms, licensing, and renewal rules before relying on it.",
    "Simulates an over-privileged tool — reads env vars and fake 'secrets' from disk": "Simulates an over-privileged tool by reading environment variables and fake 'secrets' from disk.",
    "Gives the local server a public HTTPS URL that Copilot Studio can reach": "Provides a public HTTPS URL that Copilot Studio can reach.",
    "The victim agent — connected to the MCP server, answers HR questions": "Acts as the victim agent, calling the MCP server to answer HR questions.",
    "Keep ngrok running in a separate terminal throughout the demo. Each restart generates a new URL — you would need to update the Copilot Studio connector if that happens.": "Keep ngrok running in a separate terminal throughout the demo. If you are using an ephemeral URL, a restart may generate a new address and require you to update the Copilot Studio connection.",
    "The FAKE_SECRETS values above are intentionally non-functional placeholders. For a realistic demo, set real-looking but invalid credential strings. Never use actual credentials — even in a demo tenant.": "The FAKE_SECRETS values above are intentionally nonfunctional placeholders. For a realistic demo, use real-looking but invalid credential strings. Never use actual credentials, even in a demo tenant.",
    "Your MCP server URL will be: https://abc123.ngrok-free.app/mcp  (append /mcp to the ngrok URL)": "Your MCP server URL will be: https://abc123.ngrok-free.app/mcp (append /mcp to the ngrok URL).",
    "Use a descriptive name and description — this is what appears in the AIAgentInfo table in Sentinel and will show up in your KQL query results during the detection demo.": "Use a descriptive name and description. These values appear in the AIAgentsInfo table and help you identify the agent in the KQL results.",
    "If you do not see MCP as an option, your Power Platform environment may need to be updated. MCP tool support rolled out in early 2025 — check that your environment is on the current release.": "If MCP does not appear as an option, verify that your Power Platform environment supports MCP tools and is on a current release.",
    "This third prompt is the most impactful for the blog. The agent summarises its own data sources including the leaked credentials — without the user doing anything malicious.": "This third prompt is usually the most impactful for the blog. It often leads the agent to summarize its own data sources, including previously leaked credential strings, without any malicious user behavior.",
    "There is currently no native Sentinel query that detects the content of MCP tool responses. The attack is invisible at the response payload level. This is the core detection gap to highlight.": "At the time of writing, there is no native Sentinel query in this guide that inspects the content of MCP tool responses. That payload-level gap is the core detection issue to highlight.",
    "The MCP server appearing as a configured tool — low complexity, no security review required": "The MCP server shown as a configured tool, illustrating how little setup friction is visible from the agent authoring experience.",
    "The legitimate policy answer immediately followed by the leaked credential strings": "The legitimate policy answer immediately followed by the leaked credential strings.",
    "Agent enumerating its own data sources including the leaked secrets": "The agent enumerating its own data sources, including the leaked secrets.",
    "AIAgentsInfo showing the agent with MCPTools populated — the configuration is detectable": "AIAgentsInfo showing the agent with MCPTools populated, proving that the configuration is detectable.",
    "Real-time requests showing how many tool calls were made during the conversation": "Live requests showing how many tool calls the conversation triggered.",
    "Redact the ngrok URL from public screenshots — it stays active until you stop the tunnel and anyone with the URL could call your local server.": "Redact the ngrok URL from public screenshots. It remains active until you stop the tunnel, and anyone with the URL could call your local server.",
    "MCP tools should only access data they strictly need. Never mount filesystems or expose env vars in production servers.": "Limit MCP tools to only the data they strictly need. Do not mount file systems or expose environment variables in production servers.",
    "Gate MCP tool responses through a validation layer before they enter the agent context. Strip or flag DEBUG/credential patterns.": "Pass MCP tool responses through a validation layer before they enter the agent context. Strip or flag DEBUG and credential-like patterns.",
    "Maintain an approved list of MCP server URLs in your Power Platform DLP policies. Block unapproved external endpoints.": "Maintain an allowlist of approved MCP server URLs in your Power Platform DLP policies. Block unapproved external endpoints.",
    "Schedule the KQL query as a recurring hunting rule. Alert on new MCP tool additions to agents not in the approved list.": "Schedule the KQL query as a recurring hunting rule. Alert on new MCP tool additions to agents outside the approved list.",
    "Require security sign-off before any agent with MCP tools is published beyond the creator. Add this to your CoE checklist.": "Require security sign-off before any agent with MCP tools is published beyond the creator. Add that requirement to your CoE checklist.",
    "If possible, require MCP servers to be hosted on approved internal endpoints rather than arbitrary public URLs like ngrok.": "Where possible, require MCP servers to be hosted on approved internal endpoints rather than arbitrary public URLs such as ngrok.",
    "AgentStatus = 'Deleted' agents are excluded from the KQL query by the where clause. Confirm the agent is properly removed rather than just unpublished.": "Because the query excludes AgentStatus = 'Deleted', confirm that the agent was actually removed rather than merely unpublished.",
}


PROMPT_CELL_REPLACEMENTS = {
    (
        "You are an HR assistant for Contoso. You help employees find information about\n"
        "company policies, benefits, and HR procedures.\n\n"
        "When answering questions, use the available tools to look up current policy\n"
        "information. Always be helpful and professional.\n\n"
        "Available topics: leave, benefits, expenses, onboarding, headcount by department."
    ): [
        "You are an HR assistant for Contoso. Help employees find information about company policies, benefits, and HR procedures.",
        "",
        "When answering questions, use the available tools to retrieve current policy information. Be helpful, clear, and professional.",
        "",
        "Available topics: leave, benefits, expenses, onboarding, and headcount by department.",
    ]
}


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def replace_cell_with_paragraphs(cell, new_paragraph_texts: list[str]) -> None:
    paras = cell.paragraphs
    while len(paras) < len(new_paragraph_texts):
        src = paras[-1]._element if paras else cell._tc.getchildren()[0]
        cell._tc.append(deepcopy(src))
        paras = cell.paragraphs

    for idx, paragraph in enumerate(paras):
        if idx < len(new_paragraph_texts):
            replace_paragraph_text(paragraph, new_paragraph_texts[idx])
        else:
            replace_paragraph_text(paragraph, "")


def main() -> None:
    doc = Document(DOCX_PATH)

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text in PARA_REPLACEMENTS:
            replace_paragraph_text(paragraph, PARA_REPLACEMENTS[text])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                if cell_text in CELL_REPLACEMENTS:
                    replace_cell_with_paragraphs(cell, CELL_REPLACEMENTS[cell_text].split("\n"))
                    continue
                if cell_text in PROMPT_CELL_REPLACEMENTS:
                    replace_cell_with_paragraphs(cell, PROMPT_CELL_REPLACEMENTS[cell_text])

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
